"""Multi-format document loader.

Supports: PDF, Word (.docx), Excel (.xlsx), CSV, Images, Markdown.
Extracts text, tables, and image descriptions from all formats.
Converts everything to structured text for the chunking pipeline.

For PDFs with tables: uses PyMuPDF's table extraction.
For images in PDFs: extracts and describes using Gemini Vision (if available).
For standalone images: uses Gemini Vision for description.
"""

import base64
import os
from pathlib import Path

import fitz  # PyMuPDF
import pandas as pd
from docx import Document as DocxDocument


def load_pdf(file_path: str) -> dict:
    """Extract text, tables, and image descriptions from a PDF.

    Args:
        file_path: Path to the PDF file.

    Returns:
        dict with: text (str), tables (list[str]), images (list[str]), metadata (dict)
    """
    doc = fitz.open(file_path)
    all_text = []
    all_tables = []
    all_images = []

    for page_num, page in enumerate(doc):
        # Extract text
        text = page.get_text("text")
        if text.strip():
            all_text.append(f"--- Page {page_num + 1} ---\n{text}")

        # Extract tables
        tables = page.find_tables()
        for table in tables:
            try:
                df = table.to_pandas()
                table_md = df.to_markdown(index=False)
                all_tables.append(f"[Table on page {page_num + 1}]\n{table_md}")
            except Exception:
                # Fall back to raw text if table parsing fails
                pass

        # Extract images
        images = page.get_images(full=True)
        for img_idx, img in enumerate(images):
            try:
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]

                # Try to describe image using Gemini Vision
                description = _describe_image(image_bytes, image_ext)
                if description:
                    all_images.append(
                        f"[Image on page {page_num + 1}, image {img_idx + 1}]\n{description}"
                    )
            except Exception:
                pass

    doc.close()

    combined_text = "\n\n".join(all_text)
    if all_tables:
        combined_text += "\n\n## Tables\n\n" + "\n\n".join(all_tables)
    if all_images:
        combined_text += "\n\n## Image Descriptions\n\n" + "\n\n".join(all_images)

    return {
        "text": combined_text,
        "tables": all_tables,
        "images": all_images,
        "metadata": {
            "format": "pdf",
            "pages": len(doc) if hasattr(doc, '__len__') else 0,
            "file_path": file_path,
        },
    }


def load_docx(file_path: str) -> dict:
    """Extract text and tables from a Word document.

    Args:
        file_path: Path to the .docx file.

    Returns:
        dict with: text (str), tables (list[str]), metadata (dict)
    """
    doc = DocxDocument(file_path)
    all_text = []
    all_tables = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            # Preserve heading structure
            if para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                try:
                    level = int(level)
                except ValueError:
                    level = 2
                all_text.append(f"{'#' * level} {para.text}")
            else:
                all_text.append(para.text)

    # Extract tables
    for table_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)

        if rows:
            # Convert to markdown table
            df = pd.DataFrame(rows[1:], columns=rows[0] if rows else None)
            table_md = df.to_markdown(index=False)
            all_tables.append(f"[Table {table_idx + 1}]\n{table_md}")

    combined_text = "\n\n".join(all_text)
    if all_tables:
        combined_text += "\n\n## Tables\n\n" + "\n\n".join(all_tables)

    # Extract images from docx
    all_images = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                image_data = rel.target_part.blob
                ext = rel.target_ref.split(".")[-1]
                description = _describe_image(image_data, ext)
                if description:
                    all_images.append(f"[Embedded image]\n{description}")
            except Exception:
                pass

    if all_images:
        combined_text += "\n\n## Image Descriptions\n\n" + "\n\n".join(all_images)

    return {
        "text": combined_text,
        "tables": all_tables,
        "images": all_images,
        "metadata": {
            "format": "docx",
            "paragraphs": len(doc.paragraphs),
            "file_path": file_path,
        },
    }


def load_excel(file_path: str) -> dict:
    """Extract data from Excel files (.xlsx, .xls).

    Each sheet becomes a section with the data converted to markdown table format.

    Args:
        file_path: Path to the Excel file.

    Returns:
        dict with: text (str), tables (list[str]), metadata (dict)
    """
    excel = pd.ExcelFile(file_path)
    all_text = []
    all_tables = []

    for sheet_name in excel.sheet_names:
        df = pd.read_excel(excel, sheet_name=sheet_name)

        if df.empty:
            continue

        table_md = df.to_markdown(index=False)
        section = f"## Sheet: {sheet_name}\n\n{table_md}"
        all_text.append(section)
        all_tables.append(f"[Sheet: {sheet_name}]\n{table_md}")

    return {
        "text": "\n\n".join(all_text),
        "tables": all_tables,
        "images": [],
        "metadata": {
            "format": "excel",
            "sheets": excel.sheet_names,
            "file_path": file_path,
        },
    }


def load_csv(file_path: str) -> dict:
    """Extract data from CSV files.

    Args:
        file_path: Path to the CSV file.

    Returns:
        dict with: text (str), tables (list[str]), metadata (dict)
    """
    df = pd.read_csv(file_path)
    table_md = df.to_markdown(index=False)

    return {
        "text": f"## Data\n\n{table_md}",
        "tables": [table_md],
        "images": [],
        "metadata": {
            "format": "csv",
            "rows": len(df),
            "columns": list(df.columns),
            "file_path": file_path,
        },
    }


def load_image(file_path: str) -> dict:
    """Extract description from an image using Gemini Vision.

    Args:
        file_path: Path to the image file (PNG, JPG, etc.)

    Returns:
        dict with: text (str), metadata (dict)
    """
    with open(file_path, "rb") as f:
        image_bytes = f.read()

    ext = Path(file_path).suffix.lstrip(".")
    description = _describe_image(image_bytes, ext)

    text = description if description else f"[Image: {Path(file_path).name}]"

    return {
        "text": f"## Image: {Path(file_path).name}\n\n{text}",
        "tables": [],
        "images": [text],
        "metadata": {
            "format": "image",
            "file_path": file_path,
        },
    }


def load_markdown(file_path: str) -> dict:
    """Load a markdown file (passthrough).

    Args:
        file_path: Path to the .md file.

    Returns:
        dict with: text (str), metadata (dict)
    """
    text = Path(file_path).read_text(encoding="utf-8")
    return {
        "text": text,
        "tables": [],
        "images": [],
        "metadata": {
            "format": "markdown",
            "file_path": file_path,
        },
    }


# Format registry
LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".xlsx": load_excel,
    ".xls": load_excel,
    ".csv": load_csv,
    ".png": load_image,
    ".jpg": load_image,
    ".jpeg": load_image,
    ".gif": load_image,
    ".webp": load_image,
    ".bmp": load_image,
    ".md": load_markdown,
    ".txt": load_markdown,  # Treat plain text like markdown
}


def load_document(file_path: str) -> dict:
    """Load any supported document format.

    Automatically detects format from file extension.

    Args:
        file_path: Path to the document.

    Returns:
        dict with: text (str), tables (list[str]), images (list[str]), metadata (dict)

    Raises:
        ValueError: If file format is not supported.
    """
    ext = Path(file_path).suffix.lower()

    if ext not in LOADERS:
        supported = ", ".join(sorted(LOADERS.keys()))
        raise ValueError(f"Unsupported format: {ext}. Supported: {supported}")

    return LOADERS[ext](file_path)


def load_directory(dir_path: str) -> list[dict]:
    """Load all supported documents from a directory (recursively).

    Args:
        dir_path: Path to the directory.

    Returns:
        List of loaded document dicts.
    """
    results = []
    dir_path = Path(dir_path)

    for file_path in sorted(dir_path.rglob("*")):
        if file_path.is_file() and file_path.suffix.lower() in LOADERS:
            try:
                doc = load_document(str(file_path))
                doc["source"] = file_path.stem
                doc["title"] = file_path.stem.replace("_", " ").title()
                doc["category"] = file_path.parent.name
                results.append(doc)
            except Exception as e:
                print(f"Warning: Failed to load {file_path}: {e}")

    return results


def _describe_image(image_bytes: bytes, ext: str) -> str | None:
    """Describe an image using Gemini Vision (multimodal).

    Falls back to None if Gemini is not available.
    """
    try:
        from google import genai
        from google.genai import types

        api_key = os.environ.get("GOOGLE_API_KEY", "")
        if not api_key:
            return None

        client = genai.Client(api_key=api_key)

        # Convert to base64
        b64 = base64.b64encode(image_bytes).decode("utf-8")
        mime_type = {
            "png": "image/png",
            "jpg": "image/jpeg",
            "jpeg": "image/jpeg",
            "gif": "image/gif",
            "webp": "image/webp",
            "bmp": "image/bmp",
        }.get(ext.lower(), "image/png")

        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[
                types.Part.from_bytes(data=image_bytes, mime_type=mime_type),
                "Describe this image in detail. If it contains a table, chart, or diagram, "
                "extract the data and present it as structured text. "
                "If it contains text, transcribe it.",
            ],
        )

        return response.text
    except Exception:
        return None
